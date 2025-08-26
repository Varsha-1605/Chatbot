import os
import json
import uuid
from datetime import datetime
import mimetypes
from pathlib import Path
import tempfile

# Document processing libraries
import PyPDF2
import docx
import pandas as pd
from PIL import Image
import pytesseract
import chardet

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = {
            'text': ['.txt', '.md', '.rtf'],
            'pdf': ['.pdf'],
            'word': ['.docx', '.doc'],
            'excel': ['.xlsx', '.xls', '.csv'],
            'image': ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif'],
            'json': ['.json'],
            'xml': ['.xml']
        }
        
        self.max_file_size = 10 * 1024 * 1024  # 10MB
        # With this line:
        self.upload_folder = os.environ.get('UPLOAD_FOLDER', 'uploads/documents')
        os.makedirs(self.upload_folder, exist_ok=True)
        # self.upload_folder = 'uploads/documents'
        # os.makedirs(self.upload_folder, exist_ok=True)
    
    def is_supported_format(self, filename):
        """Check if file format is supported"""
        file_ext = Path(filename).suffix.lower()
        for format_type, extensions in self.supported_formats.items():
            if file_ext in extensions:
                return True, format_type
        return False, None
    
    def validate_file(self, file_data, filename):
        """Validate uploaded file"""
        errors = []
        
        # Check file size
        if len(file_data) > self.max_file_size:
            errors.append(f"File size exceeds {self.max_file_size // 1024 // 1024}MB limit")
        
        # Check file format
        is_supported, file_type = self.is_supported_format(filename)
        if not is_supported:
            errors.append("File format not supported")
        
        return errors, file_type
    
    def save_uploaded_file(self, file_data, filename, user_id):
        """Save uploaded file and return file info"""
        file_id = str(uuid.uuid4())
        file_ext = Path(filename).suffix.lower()
        safe_filename = f"{file_id}{file_ext}"
        file_path = os.path.join(self.upload_folder, safe_filename)
        
        # Save file
        with open(file_path, 'wb') as f:
            f.write(file_data)
        
        # Create file info
        file_info = {
            'file_id': file_id,
            'original_filename': filename,
            'saved_filename': safe_filename,
            'file_path': file_path,
            'file_size': len(file_data),
            'uploaded_at': datetime.now().isoformat(),
            'user_id': user_id,
            'processed': False
        }
        
        return file_info
    
    def extract_text_from_file(self, file_path, file_type):
        """Extract text content from various file types"""
        try:
            text_content = ""
            
            if file_type == 'text':
                text_content = self._extract_from_text(file_path)
            elif file_type == 'pdf':
                text_content = self._extract_from_pdf(file_path)
            elif file_type == 'word':
                text_content = self._extract_from_word(file_path)
            elif file_type == 'excel':
                text_content = self._extract_from_excel(file_path)
            elif file_type == 'image':
                text_content = self._extract_from_image(file_path)
            elif file_type == 'json':
                text_content = self._extract_from_json(file_path)
            elif file_type == 'xml':
                text_content = self._extract_from_xml(file_path)
            
            return text_content.strip(), None
            
        except Exception as e:
            return "", str(e)
    
    def _extract_from_text(self, file_path):
        """Extract text from plain text files"""
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
            
            # Read with detected encoding
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except:
            # Fallback to utf-8 with error handling
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
    
    def _extract_from_pdf(self, file_path):
        """Extract text from PDF files"""
        text = ""
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            raise Exception(f"PDF extraction error: {str(e)}")
        return text
    
    def _extract_from_word(self, file_path):
        """Extract text from Word documents"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            raise Exception(f"Word document extraction error: {str(e)}")
    
    def _extract_from_excel(self, file_path):
        """Extract text from Excel/CSV files"""
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.csv':
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path, sheet_name=None)
                
                # If multiple sheets, combine them
                if isinstance(df, dict):
                    combined_text = ""
                    for sheet_name, sheet_df in df.items():
                        combined_text += f"Sheet: {sheet_name}\n"
                        combined_text += sheet_df.to_string(index=False) + "\n\n"
                    return combined_text
                else:
                    df = df
            
            return df.to_string(index=False)
            
        except Exception as e:
            raise Exception(f"Excel/CSV extraction error: {str(e)}")
    
    def _extract_from_image(self, file_path):
        """Extract text from images using OCR"""
        try:
            # Check if tesseract is available
            try:
                import pytesseract
            except ImportError:
                raise Exception("OCR functionality requires pytesseract library")
            
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            raise Exception(f"Image OCR extraction error: {str(e)}")
    
    def _extract_from_json(self, file_path):
        """Extract and format JSON content"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert JSON to readable text format
            def json_to_text(obj, indent=0):
                text = ""
                spaces = "  " * indent
                
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        if isinstance(value, (dict, list)):
                            text += f"{spaces}{key}:\n"
                            text += json_to_text(value, indent + 1)
                        else:
                            text += f"{spaces}{key}: {value}\n"
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        text += f"{spaces}[{i}]:\n"
                        text += json_to_text(item, indent + 1)
                else:
                    text += f"{spaces}{obj}\n"
                
                return text
            
            return json_to_text(data)
            
        except Exception as e:
            raise Exception(f"JSON extraction error: {str(e)}")
    
    def _extract_from_xml(self, file_path):
        """Extract text from XML files"""
        try:
            import xml.etree.ElementTree as ET
            
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            def xml_to_text(element, indent=0):
                text = ""
                spaces = "  " * indent
                
                # Add element tag and attributes
                if element.attrib:
                    attrs = " ".join([f"{k}={v}" for k, v in element.attrib.items()])
                    text += f"{spaces}<{element.tag} {attrs}>\n"
                else:
                    text += f"{spaces}<{element.tag}>\n"
                
                # Add text content
                if element.text and element.text.strip():
                    text += f"{spaces}  {element.text.strip()}\n"
                
                # Process children
                for child in element:
                    text += xml_to_text(child, indent + 1)
                
                return text
            
            return xml_to_text(root)
            
        except Exception as e:
            raise Exception(f"XML extraction error: {str(e)}")
    
    def process_document(self, file_info):
        """Process uploaded document and extract content"""
        try:
            # Determine file type
            _, file_type = self.is_supported_format(file_info['original_filename'])
            
            # Extract text content
            content, error = self.extract_text_from_file(file_info['file_path'], file_type)
            
            if error:
                return {
                    'success': False,
                    'error': error,
                    'content': None
                }
            
            # Create document summary
            word_count = len(content.split())
            char_count = len(content)
            
            processing_result = {
                'success': True,
                'content': content,
                'metadata': {
                    'file_type': file_type,
                    'word_count': word_count,
                    'char_count': char_count,
                    'processed_at': datetime.now().isoformat()
                },
                'summary': self._create_content_summary(content, file_info['original_filename'])
            }
            
            # Update file info
            file_info['processed'] = True
            file_info['processing_result'] = processing_result
            
            return processing_result
            
        except Exception as e:
            return {
                'success': False,
                'error': f"Processing error: {str(e)}",
                'content': None
            }
    
    def _create_content_summary(self, content, filename):
        """Create a brief summary of the document content"""
        lines = content.split('\n')
        non_empty_lines = [line.strip() for line in lines if line.strip()]
        
        summary = {
            'filename': filename,
            'total_lines': len(lines),
            'non_empty_lines': len(non_empty_lines),
            'first_few_lines': non_empty_lines[:5] if non_empty_lines else [],
            'word_count': len(content.split()),
            'char_count': len(content)
        }
        
        return summary
    
    def cleanup_old_files(self, max_age_days=7):
        """Clean up old uploaded files"""
        try:
            current_time = datetime.now()
            for filename in os.listdir(self.upload_folder):
                file_path = os.path.join(self.upload_folder, filename)
                if os.path.isfile(file_path):
                    file_time = datetime.fromtimestamp(os.path.getctime(file_path))
                    age_days = (current_time - file_time).days
                    
                    if age_days > max_age_days:
                        os.remove(file_path)
        except Exception as e:
            print(f"Cleanup error: {str(e)}")